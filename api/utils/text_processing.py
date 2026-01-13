"""텍스트 처리 유틸리티 - 고급 청킹 및 문서 파싱"""

import re
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path

logger = logging.getLogger(__name__)


def clean_text(text: str) -> str:
    """
    텍스트 정제 및 정규화
    
    Args:
        text: 입력 텍스트
        
    Returns:
        정제된 텍스트
    """
    # 여러 공백을 하나로
    text = re.sub(r'\s+', ' ', text)
    
    # 특수문자 제거 (구두점은 유지)
    text = re.sub(r'[^\w\s.,!?;:\'-]', '', text)
    
    return text.strip()


def chunk_text_advanced(
    text: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    separator: str = "\n\n"
) -> List[Dict[str, Any]]:
    """
    고급 텍스트 청킹 - 의미 단위 보존
    
    LangChain의 RecursiveCharacterTextSplitter 스타일 구현
    
    Args:
        text: 입력 텍스트
        chunk_size: 청크 최대 크기 (문자 수)
        chunk_overlap: 청크 간 겹침 크기
        separator: 우선 분할 구분자
        
    Returns:
        청크 리스트 (메타데이터 포함)
        [
            {
                "content": "청크 내용",
                "metadata": {
                    "chunk_index": 0,
                    "char_start": 0,
                    "char_end": 1000,
                    "word_count": 150
                }
            }
        ]
    """
    if len(text) <= chunk_size:
        return [{
            "content": text,
            "metadata": {
                "chunk_index": 0,
                "char_start": 0,
                "char_end": len(text),
                "word_count": len(text.split())
            }
        }]
    
    chunks = []
    
    # 분할 우선순위 (의미 단위 보존)
    separators = [
        "\n\n",      # 단락
        "\n",        # 줄바꿈
        ". ",        # 문장
        "! ",        # 느낌표
        "? ",        # 물음표
        "; ",        # 세미콜론
        ", ",        # 쉼표
        " ",         # 공백
        ""           # 문자 단위
    ]
    
    def _split_text(text: str, separators: List[str]) -> List[str]:
        """재귀적으로 텍스트 분할"""
        if not separators:
            # 마지막 구분자: 문자 단위
            return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size - chunk_overlap)]
        
        separator = separators[0]
        splits = text.split(separator) if separator else [text]
        
        result = []
        current_chunk = ""
        
        for split in splits:
            # 분할 조각이 너무 크면 다음 구분자로 재귀
            if len(split) > chunk_size:
                if current_chunk:
                    result.append(current_chunk)
                    current_chunk = ""
                result.extend(_split_text(split, separators[1:]))
            else:
                # 현재 청크에 추가 가능한지 확인
                if len(current_chunk) + len(split) + len(separator) <= chunk_size:
                    current_chunk += (separator if current_chunk else "") + split
                else:
                    if current_chunk:
                        result.append(current_chunk)
                    current_chunk = split
        
        if current_chunk:
            result.append(current_chunk)
        
        return result
    
    raw_chunks = _split_text(text, separators)
    
    # 겹침 처리 및 메타데이터 추가
    char_position = 0
    for idx, chunk_text in enumerate(raw_chunks):
        chunk_text = chunk_text.strip()
        if not chunk_text:
            continue
        
        # 이전 청크와 겹침 추가
        if idx > 0 and chunk_overlap > 0:
            prev_chunk = raw_chunks[idx - 1]
            overlap_text = prev_chunk[-chunk_overlap:].strip()
            chunk_text = overlap_text + " " + chunk_text
        
        chunks.append({
            "content": chunk_text,
            "metadata": {
                "chunk_index": idx,
                "char_start": char_position,
                "char_end": char_position + len(chunk_text),
                "word_count": len(chunk_text.split()),
                "char_count": len(chunk_text)
            }
        })
        
        char_position += len(chunk_text)
    
    logger.info(f"Split text into {len(chunks)} chunks (size={chunk_size}, overlap={chunk_overlap})")
    return chunks


def chunk_text_by_sentences(
    text: str,
    max_sentences: int = 5,
    min_chunk_size: int = 100
) -> List[Dict[str, Any]]:
    """
    문장 단위로 청킹 (더 자연스러운 분할)
    
    Args:
        text: 입력 텍스트
        max_sentences: 청크당 최대 문장 수
        min_chunk_size: 최소 청크 크기 (문자 수)
        
    Returns:
        청크 리스트
    """
    # 문장 분리 (간단한 정규식 - 실제로는 NLTK 등 사용 권장)
    sentences = re.split(r'(?<=[.!?])\s+', text)
    
    chunks = []
    current_chunk = []
    current_size = 0
    
    for sentence in sentences:
        current_chunk.append(sentence)
        current_size += len(sentence)
        
        # 청크 조건 만족 시 저장
        if len(current_chunk) >= max_sentences or current_size >= min_chunk_size * max_sentences:
            chunk_text = " ".join(current_chunk)
            chunks.append({
                "content": chunk_text,
                "metadata": {
                    "chunk_index": len(chunks),
                    "sentence_count": len(current_chunk),
                    "char_count": len(chunk_text)
                }
            })
            current_chunk = []
            current_size = 0
    
    # 남은 문장 처리
    if current_chunk:
        chunk_text = " ".join(current_chunk)
        chunks.append({
            "content": chunk_text,
            "metadata": {
                "chunk_index": len(chunks),
                "sentence_count": len(current_chunk),
                "char_count": len(chunk_text)
            }
        })
    
    return chunks


def extract_text_from_pdf(file_path: str) -> str:
    """
    PDF 파일에서 텍스트 추출
    
    Args:
        file_path: PDF 파일 경로
        
    Returns:
        추출된 텍스트
    """
    try:
        import PyPDF2
        
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n\n"
        
        logger.info(f"Extracted text from PDF: {num_pages} pages, {len(text)} characters")
        return text.strip()
    
    except Exception as e:
        logger.error(f"PDF text extraction failed: {e}")
        raise


def extract_text_from_docx(file_path: str) -> str:
    """
    DOCX 파일에서 텍스트 추출
    
    Args:
        file_path: DOCX 파일 경로
        
    Returns:
        추출된 텍스트
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        # 단락 추출
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # 표 추출
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text.append(" | ".join(row_text))
        
        # 결합
        text = "\n\n".join(paragraphs)
        if table_text:
            text += "\n\n" + "\n".join(table_text)
        
        logger.info(f"Extracted text from DOCX: {len(doc.paragraphs)} paragraphs, {len(text)} characters")
        return text.strip()
    
    except Exception as e:
        logger.error(f"DOCX text extraction failed: {e}")
        raise


def extract_text_from_file(file_path: str) -> str:
    """
    파일 확장자에 따라 자동으로 텍스트 추출
    
    Args:
        file_path: 파일 경로
        
    Returns:
        추출된 텍스트
    """
    file_extension = Path(file_path).suffix.lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif file_extension in ['.txt', '.md']:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")


def extract_keywords(text: str, max_keywords: int = 10) -> List[str]:
    """
    텍스트에서 키워드 추출 (빈도 기반)
    
    Args:
        text: 입력 텍스트
        max_keywords: 최대 키워드 수
        
    Returns:
        키워드 리스트
    """
    # 불용어
    stop_words = {
        'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
        'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'were', 'been',
        'be', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would',
        'should', 'could', 'may', 'might', 'can', 'this', 'that', 'these',
        'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they', 'them', 'their'
    }
    
    # 토큰화 및 카운트
    words = re.findall(r'\b\w+\b', text.lower())
    word_freq = {}
    
    for word in words:
        if len(word) > 3 and word not in stop_words:
            word_freq[word] = word_freq.get(word, 0) + 1
    
    # 빈도순 정렬
    keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
    
    return [word for word, freq in keywords[:max_keywords]]


def truncate_text(text: str, max_length: int = 1000, suffix: str = "...") -> str:
    """
    텍스트를 최대 길이로 자르기
    
    Args:
        text: 입력 텍스트
        max_length: 최대 길이
        suffix: 잘렸을 때 추가할 접미사
        
    Returns:
        잘린 텍스트
    """
    if len(text) <= max_length:
        return text
    
    return text[:max_length - len(suffix)] + suffix


def estimate_tokens(text: str) -> int:
    """
    대략적인 토큰 수 추정 (1 토큰 ≈ 4 문자)
    
    Args:
        text: 입력 텍스트
        
    Returns:
        추정 토큰 수
    """
    return len(text) // 4


def chunk_by_tokens(
    text: str,
    max_tokens: int = 512,
    overlap_tokens: int = 50
) -> List[Dict[str, Any]]:
    """
    토큰 기반 청킹 (LLM 컨텍스트 윈도우에 맞춤)
    
    Args:
        text: 입력 텍스트
        max_tokens: 청크당 최대 토큰 수
        overlap_tokens: 겹침 토큰 수
        
    Returns:
        청크 리스트
    """
    # 문자 수로 변환 (1 토큰 ≈ 4 문자)
    chunk_size = max_tokens * 4
    overlap_size = overlap_tokens * 4
    
    return chunk_text_advanced(text, chunk_size, overlap_size)


def smart_chunk(
    text: str,
    file_type: str = "text",
    max_chunk_size: int = 1000
) -> List[Dict[str, Any]]:
    """
    파일 타입과 크기에 따라 최적의 청킹 방법 선택
    
    Args:
        text: 입력 텍스트
        file_type: 파일 타입 ("text", "pdf", "docx")
        max_chunk_size: 최대 청크 크기
        
    Returns:
        청크 리스트
    """
    text_length = len(text)
    
    # 짧은 텍스트는 청킹 불필요
    if text_length <= max_chunk_size:
        return [{
            "content": text,
            "metadata": {
                "chunk_index": 0,
                "char_count": text_length,
                "chunking_method": "none"
            }
        }]
    
    # PDF/긴 문서: 토큰 기반 청킹
    if file_type in ["pdf", "docx"] or text_length > 10000:
        logger.info(f"Using token-based chunking for {file_type}")
        chunks = chunk_by_tokens(text, max_tokens=max_chunk_size // 4)
        for chunk in chunks:
            chunk["metadata"]["chunking_method"] = "token_based"
        return chunks
    
    # 일반 텍스트: 문장 기반 청킹
    else:
        logger.info(f"Using sentence-based chunking")
        chunks = chunk_text_by_sentences(text, max_sentences=5)
        for chunk in chunks:
            chunk["metadata"]["chunking_method"] = "sentence_based"
        return chunks
